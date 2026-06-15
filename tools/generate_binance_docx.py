from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "output" / "docx"
PDF_DATA_PATH = ROOT / "tools" / "generate_resume_pdfs.py"
OUTPUT_PATH = OUTPUT_DIR / "jay-wu-binance-en.docx"


COLORS = {
    "ink": "172A3A",
    "blue": "1F4D78",
    "accent": "2E74B5",
    "muted": "5B6670",
    "line": "D6E1EC",
    "fill": "F4F7FB",
    "label_fill": "E8EEF5",
}


def load_spec():
    spec = spec_from_file_location("resume_pdf_data", PDF_DATA_PATH)
    module = module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.DATA["jay-wu-binance-en.pdf"]


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D6E1EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:{}".format(margin)))
        if node is None:
            node = OxmlElement("w:{}".format(margin))
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    return run


def add_hyperlink(paragraph, text, url, size=8.4):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["accent"])
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    run_props.append(font_size)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_props)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_paragraph(paragraph, before=0, after=4, line=1.15, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=10, after=4, line=1.0)
    add_run(p, text.upper(), bold=True, size=10.5, color=COLORS["blue"])
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), COLORS["line"])
    border.append(bottom)
    p_pr.append(border)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=2, line=1.12)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    add_run(p, text, size=9.2, color=COLORS["ink"])


def add_meta_line(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=2, line=1.0)
    add_run(p, text, size=9.1, color=COLORS["muted"])


def build_docx():
    data = load_spec()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    style_paragraph(header, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(header, "Jay Wu | Binance-focused CV", size=8, color=COLORS["muted"])

    footer = section.footer.paragraphs[0]
    style_paragraph(footer, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(footer, "Prepared for Binance roles | English version for interview use", size=8, color=COLORS["muted"])

    title = doc.add_paragraph()
    style_paragraph(title, before=0, after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(title, data["name"], bold=True, size=20, color=COLORS["ink"])

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, before=0, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(subtitle, data["title"], bold=True, size=11, color=COLORS["accent"])

    contact = doc.add_paragraph()
    style_paragraph(contact, before=0, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(contact, "Email: dan40912@gmail.com | ", size=8.4, color=COLORS["muted"])
    add_hyperlink(
        contact,
        "LinkedIn Profile",
        "https://www.linkedin.com/in/%E6%96%B0%E6%8D%B7-%E5%90%B3-630779107/",
    )
    add_run(contact, " | ", size=8.4, color=COLORS["muted"])
    add_hyperlink(contact, "Medium", "https://medium.com/@dan40912")

    add_section_heading(doc, "Profile")
    summary = doc.add_paragraph()
    style_paragraph(summary, before=0, after=5, line=1.15)
    add_run(summary, data["summary"], size=9.5, color=COLORS["ink"])

    add_section_heading(doc, data["labels"]["proof"])
    for item in data["proof"]:
        add_bullet(doc, item)

    add_section_heading(doc, data["labels"]["experience"])
    for item in data["experience"]:
        p = doc.add_paragraph()
        style_paragraph(p, before=3, after=0, line=1.0)
        add_run(p, item["title"], bold=True, size=10, color=COLORS["ink"])
        add_meta_line(doc, item["meta"])
        for bullet in item["bullets"]:
            add_bullet(doc, bullet)

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_section_heading(doc, data["labels"]["projects"])
    for item in data["projects"]:
        p = doc.add_paragraph()
        style_paragraph(p, before=4, after=0, line=1.0)
        add_run(p, item["title"], bold=True, size=10, color=COLORS["ink"])
        add_meta_line(doc, item["meta"])
        summary_p = doc.add_paragraph()
        style_paragraph(summary_p, before=0, after=1, line=1.12)
        add_run(summary_p, item["summary"], size=9.1, color=COLORS["ink"])
        for bullet in item["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, data["labels"]["skills"])
    skills = doc.add_table(rows=2, cols=2)
    set_table_width(skills, [3.35, 3.35])
    for row in skills.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
            set_cell_shading(cell, "FFFFFF")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for idx, (label, value) in enumerate(data["skills"]):
        cell = skills.rows[idx // 2].cells[idx % 2]
        heading = cell.paragraphs[0]
        style_paragraph(heading, before=0, after=2, line=1.0)
        add_run(heading, label, bold=True, size=8.9, color=COLORS["blue"])
        for skill in [part.strip() for part in value.split(";") if part.strip()]:
            p = cell.add_paragraph(style="List Bullet")
            style_paragraph(p, before=0, after=0, line=1.05)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.10)
            add_run(p, skill, size=8.0, color=COLORS["ink"])

    add_section_heading(doc, "Note")
    note = doc.add_paragraph()
    style_paragraph(note, before=0, after=0, line=1.1)
    add_run(note, data["note"], size=8.7, color=COLORS["muted"])

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_docx()
